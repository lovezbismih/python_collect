from selenium import webdriver
from selenium.webdriver.common.by import By
from time import sleep
from pywinauto import Application
from docx import Document
import os
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
 
file_path = r'D:\c\a.docx'  # 使用原始字符串来避免转义问题
 
# 打印文件路径，以便检查
print(file_path)
 
try:
    doc = Document(file_path)
    print('文档加载成功。')
    # 你的其他操作
except FileNotFoundError:
    print("文件未找到，请检查文件路径是否正确。")
folder_path = r'D:\a\a'  # 使用原始字符串，确保路径正确
 
'''try:
    files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
    # 你的其他操作
except FileNotFoundError:
    print(f"找不到指定的路径：{folder_path}")'''
driver = webdriver.Edge()
driver.maximize_window()
driver.get('https://cloud.tencent.com/login')
sleep(10)  # 等待登录完成，必须在10秒内完成，否则会报错
 
files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
 
n = 0
for file in files:
    i = n % 2
    n = n + 1
    driver.get('https://cloud.tencent.com/product/ocr')
    driver.execute_script("window.scrollBy(0, 400);")
    WebDriverWait(driver, 20).until(EC.presence_of_element_located(
        (By.XPATH, '/html/body/div[2]/div/div/div[2]/div/div/div/div/div[1]/div[2]/div/div[2]'))).click()
    sleep(0.1)
    anniu = driver.find_element(By.XPATH,
                                '//*[@id="demo-app"]/div/div/div/div[2]/div[2]/div/div[1]/div[2]/div[1]/div[1]')
    driver.execute_script("arguments[0].scrollIntoView(true);", anniu)
    anniu.click()
    sleep(0.5)
    app = Application().connect(title='打开')
    window = app.window(title='打开')
    window["文件名(&N):Edit"].set_text(os.path.join(folder_path, file))
    sleep(0.2)
    window["文件名(&N):Edit"].type_keys("{VK_RETURN}")
    sleep(3)
    table = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.TAG_NAME, 'tbody')))
    rows = table.find_elements(By.TAG_NAME, "tr")
    # 定义文件保存的路径
    file_path = 'd:/c/a.docx'
 
    # 获取文件所在的目录
    dir_path = os.path.dirname(file_path)
 
    # 如果目录不存在，则创建它
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    doc = Document(file_path)
    # 遍历每一行，获取单元格中的文本
    if i == 0:
        doc.add_paragraph('左±×·℃÷≤≥≠≈★√∵∴∠πμρηΔΩ°′″________[   ]①②③④⑤')
    '''else:
        doc.add_paragraph('右±×·℃÷≤≥≠≈★√∵∴∠πμρηΔΩ°′″________[   ]①②③④⑤')'''
    for row in rows:
        cells = row.find_elements(By.TAG_NAME, "td")
        for cell in cells:
            doc.add_paragraph(cell.text)
    # 保存文档
    doc.save(file_path)